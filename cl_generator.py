from docx import Document
import argparse
import os
import datetime

def set_para_data(out_doc, in_para, modified_text):
    out_para = out_doc.add_paragraph(modified_text)
    para_format_attr = ['alignment', 'first_line_indent', 'keep_together', 'keep_with_next', 'left_indent', 'line_spacing', 'line_spacing_rule', 
    'page_break_before', 'right_indent', 'space_after', 'space_before', 'widow_control']
    for attr in para_format_attr:
        """
        Lazy way to copy all attributes from source to destination. Instead of writing 
        out line-by-line assignment of each attribute, we use python built-in methods
        getattr to get the value of an attribute by name and setattr to set the attribute
        by name. Find out more more at https://docs.python.org/3.6/library/functions.html.
        """
        setattr(out_para.paragraph_format, attr, getattr(in_para.paragraph_format, attr))
    out_para.style = in_para.style
    out_para.text = modified_text
    return out_para

def generate(args):
    company_name, job_position = args.company_name, args.job_position
    output_path = os.path.join(os.getcwd(), "output")
    current_time = datetime.datetime.now().strftime("%m.%d.%H.%M")
    
    document = Document("cl_template.docx")
    if not os.path.exists(output_path):
        os.mkdir(output_path)
    out_doc = Document()
    for para in document.paragraphs:
        modified_text = para.text.format(**{'company_name':company_name, 'job_position':job_position})
        out_para = set_para_data(out_doc, para, modified_text)
    out_doc.save(f"{output_path}/{company_name}_cover_letter_{job_position}_{current_time}.docx")
    print("Generated Successfully :)")
    
available_args = {
    'jp': 'job_position',
    'cn': 'company_name'
}
parser = argparse.ArgumentParser(description='Input')
for short_name, long_name in available_args.items():
    parser.add_argument(f'-{short_name}', f'--{long_name}', dest=f'{long_name}',
                        help=f'Desire {long_name}')

args = parser.parse_args()
generate(args)

